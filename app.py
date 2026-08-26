import streamlit as st
from pypdf import PdfWriter, PdfReader
from datetime import datetime
import re
import random
import string
import io
import pandas as pd

# ======================
# Page Config
# ======================
st.set_page_config(
    page_title="PDF Tools",
    page_icon="📄",
    layout="centered"
)

# ======================
# Helper Functions
# ======================
def is_valid_email(email: str) -> bool:
    pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
    return bool(re.match(pattern, email.strip()))


def generate_verification_code(length=6):
    return "".join(random.choices(string.digits, k=length))


def merge_pdfs(uploaded_files):
    writer = PdfWriter()
    for uploaded_file in uploaded_files:
        try:
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                writer.add_page(page)
        except Exception as e:
            st.error(f"Error reading {uploaded_file.name}: {e}")
            return None
    return writer


def extract_pages(pdf_bytes, selected_pages):
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page_num in selected_pages:
        idx = page_num - 1
        if 0 <= idx < len(reader.pages):
            writer.add_page(reader.pages[idx])
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def rotate_pages(pdf_bytes, selected_pages, angle):
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    selected_set = set(selected_pages)
    for i, page in enumerate(reader.pages):
        page_num = i + 1
        if page_num in selected_set:
            writer.add_page(page.rotate(angle))
        else:
            writer.add_page(page)
    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def has_extractable_text(pdf_bytes: bytes, min_chars: int = 30) -> bool:
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        total_text = ""
        for page in reader.pages[:5]:
            text = page.extract_text() or ""
            total_text += text
            if len(total_text.strip()) >= min_chars:
                return True
        return len(total_text.strip()) >= min_chars
    except Exception:
        return False


def pdf_to_word_with_pdf2docx(pdf_bytes: bytes) -> bytes:
    from pdf2docx import Converter
    docx_stream = io.BytesIO()
    cv = Converter(stream=pdf_bytes)
    try:
        cv.convert(docx_stream)
    finally:
        cv.close()
    docx_stream.seek(0)
    return docx_stream.getvalue()


def pdf_to_word_with_ocr(pdf_bytes: bytes, lang: str = "eng") -> bytes:
    import pypdfium2 as pdfium
    import pytesseract
    from docx import Document
    from docx.shared import Pt

    pdf = pdfium.PdfDocument(pdf_bytes)
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    for i in range(len(pdf)):
        page = pdf[i]
        bitmap = page.render(scale=2.0)
        pil_image = bitmap.to_pil()
        text = pytesseract.image_to_string(pil_image, lang=lang)

        if i > 0:
            doc.add_page_break()
        doc.add_paragraph(f"--- Page {i + 1} ---")
        doc.add_paragraph(text.strip() if text.strip() else "[No text detected on this page]")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def pdf_to_word(pdf_bytes: bytes, force_ocr: bool = False, ocr_lang: str = "eng") -> bytes:
    try:
        from pdf2docx import Converter
    except ImportError:
        raise ImportError("Missing package: pdf2docx\nPlease install it with: pip install pdf2docx")

    if force_ocr or not has_extractable_text(pdf_bytes):
        try:
            return pdf_to_word_with_ocr(pdf_bytes, lang=ocr_lang)
        except Exception as e:
            raise RuntimeError(f"OCR conversion failed: {e}")
    else:
        return pdf_to_word_with_pdf2docx(pdf_bytes)


def preview_pdf(pdf_bytes: bytes, key: str = "pdf_preview", max_pages: int = 8):
    if not pdf_bytes:
        st.warning("No PDF data to preview.")
        return

    missing = []
    try:
        import pypdfium2 as pdfium
    except ImportError:
        missing.append("pypdfium2")
    try:
        from PIL import Image  # noqa: F401
    except ImportError:
        missing.append("Pillow")

    if missing:
        st.error(f"❌ Missing package(s): {', '.join(missing)}")
        st.code(f"pip install {' '.join(missing)}", language="bash")
        return

    try:
        pdf = pdfium.PdfDocument(pdf_bytes)
        n_pages = len(pdf)
        st.caption(f"📄 Preview — showing {min(n_pages, max_pages)} of {n_pages} page(s)")
        for i in range(min(n_pages, max_pages)):
            page = pdf[i]
            bitmap = page.render(scale=1.5)
            pil_image = bitmap.to_pil()
            st.image(pil_image, caption=f"Page {i + 1}", use_container_width=True)
        if n_pages > max_pages:
            st.info(f"Only the first {max_pages} pages are shown for performance.")
    except Exception as e:
        st.error(f"Could not render PDF preview: {e}")


def preview_and_edit_pdf(pdf_bytes: bytes, key: str = "pdf_edit", max_pages: int = 20):
    if not pdf_bytes:
        st.warning("No PDF data to preview/edit.")
        return None

    try:
        import pypdfium2 as pdfium
        from PIL import Image
    except ImportError as e:
        st.error(f"Missing package: {e}")
        st.code("pip install pypdfium2 Pillow")
        return None

    try:
        pdf = pdfium.PdfDocument(pdf_bytes)
        n_pages = len(pdf)
    except Exception as e:
        st.error(f"Cannot open PDF: {e}")
        return None

    st.caption(f"📄 Preview & Edit — {n_pages} page(s)")

    st.markdown("**1. Select pages to keep**")
    page_options = list(range(1, n_pages + 1))
    selected_pages = st.multiselect(
        "Pages to include in the final PDF (order will be preserved)",
        options=page_options,
        default=page_options,
        key=f"{key}_select_pages"
    )

    if not selected_pages:
        st.warning("Please select at least one page.")
        return None

    st.markdown("**2. Rotate pages (optional)**")
    rotate_choice = st.selectbox(
        "Rotate selected pages by",
        options=[0, 90, 180, 270],
        format_func=lambda x: "No rotation" if x == 0 else f"{x}° clockwise",
        key=f"{key}_rotate"
    )

    st.markdown("**3. Preview of selected pages**")
    cols = st.columns(2)

    for idx, page_num in enumerate(selected_pages[:max_pages]):
        page = pdf[page_num - 1]
        bitmap = page.render(scale=1.3)
        pil_image = bitmap.to_pil()

        if rotate_choice != 0:
            pil_image = pil_image.rotate(-rotate_choice, expand=True)

        with cols[idx % 2]:
            st.image(
                pil_image,
                caption=f"Page {page_num}" + (f" (rotated {rotate_choice}°)" if rotate_choice else ""),
                use_container_width=True
            )

    if len(selected_pages) > max_pages:
        st.info(f"Showing first {max_pages} of {len(selected_pages)} selected pages.")

    st.markdown("---")
    if st.button("✅ Apply edits & create new PDF", type="primary", key=f"{key}_apply"):
        with st.spinner("Building edited PDF..."):
            try:
                reader = PdfReader(io.BytesIO(pdf_bytes))
                writer = PdfWriter()

                for page_num in selected_pages:
                    page = reader.pages[page_num - 1]
                    if rotate_choice != 0:
                        page = page.rotate(rotate_choice)
                    writer.add_page(page)

                buffer = io.BytesIO()
                writer.write(buffer)
                buffer.seek(0)
                edited_bytes = buffer.getvalue()

                st.success(
                    f"✅ Edited PDF created! "
                    f"{len(selected_pages)} page(s)"
                    + (f", rotated {rotate_choice}°" if rotate_choice else "")
                )
                return edited_bytes

            except Exception as e:
                st.error(f"Failed to create edited PDF: {e}")
                return None

    return None


def extract_text_from_pdf(pdf_bytes: bytes, use_ocr: bool = False, ocr_lang: str = "eng") -> str:
    """Extract all text from PDF (with optional OCR)."""
    if use_ocr or not has_extractable_text(pdf_bytes):
        try:
            import pypdfium2 as pdfium
            import pytesseract
            pdf = pdfium.PdfDocument(pdf_bytes)
            texts = []
            for i in range(len(pdf)):
                page = pdf[i]
                bitmap = page.render(scale=2.0)
                pil_image = bitmap.to_pil()
                text = pytesseract.image_to_string(pil_image, lang=ocr_lang)
                texts.append(text)
            return "\n".join(texts)
        except Exception as e:
            return f"[OCR failed: {e}]"
    else:
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            texts = []
            for page in reader.pages:
                texts.append(page.extract_text() or "")
            return "\n".join(texts)
        except Exception as e:
            return f"[Text extraction failed: {e}]"


def extract_agreement_info(pdf_bytes: bytes, use_ocr: bool = False, ocr_lang: str = "eng") -> dict:
    """
    Extract common agreement fields from a PDF:
    - Agreement / Effective Date
    - Party Name & Address
    - Signor Name & Title
    """
    text = extract_text_from_pdf(pdf_bytes, use_ocr=use_ocr, ocr_lang=ocr_lang)
    result = {
        "agreement_date": None,
        "parties": [],
        "signors": [],
        "raw_text_preview": text[:3000] + ("..." if len(text) > 3000 else "")
    }

    # 1. Agreement / Effective Date
    date_patterns = [
        r"(?:agreement|effective|commencement|dated|date of this agreement|as of)[:\s]*([0-9]{1,2}[\/\-\.][0-9]{1,2}[\/\-\.][0-9]{2,4})",
        r"(?:agreement|effective|commencement|dated|date of this agreement|as of)[:\s]*([0-9]{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+[0-9]{2,4})",
        r"(?:agreement|effective|commencement|dated|date of this agreement|as of)[:\s]*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+[0-9]{1,2},?\s+[0-9]{2,4})",
        r"dated\s+this\s+([0-9]{1,2}(?:st|nd|rd|th)?\s+day\s+of\s+\w+\s+[0-9]{4})",
        r"([0-9]{1,2}[\/\-\.][0-9]{1,2}[\/\-\.][0-9]{2,4})",
    ]
    for pattern in date_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result["agreement_date"] = match.group(1).strip()
            break

    # 2. Parties (Name + Address)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    potential_parties = []
    for i, line in enumerate(lines):
        if any(keyword in line.lower() for keyword in [
            "limited", "ltd", "inc", "llc", "corporation", "company",
            "party a", "party b", "the company", "the client"
        ]):
            address = ""
            if i + 1 < len(lines):
                address = lines[i + 1]
            if i + 2 < len(lines) and len(lines[i + 2]) > 8:
                address += ", " + lines[i + 2]
            potential_parties.append({
                "name": line[:120],
                "address": address[:180]
            })
    # Deduplicate roughly
    seen_names = set()
    unique_parties = []
    for p in potential_parties:
        key = p["name"].lower()[:40]
        if key not in seen_names:
            seen_names.add(key)
            unique_parties.append(p)
    result["parties"] = unique_parties[:5]

    # 3. Signor Name & Title
    signor_patterns = [
        r"(?:signed by|signature|for and on behalf of|authorised signatory|authorized signatory)[:\s]*([A-Z][a-zA-Z\s\.]{2,50})",
        r"(?:name)[:\s]*([A-Z][a-zA-Z\s\.]{2,50})\s*(?:title|position|designation)[:\s]*([A-Za-z\s]{2,50})",
        r"(?:title|position|designation)[:\s]*([A-Za-z\s]{2,50})",
    ]
    signors = []
    for pattern in signor_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for m in matches:
            if isinstance(m, tuple):
                name = m[0].strip() if m[0] else None
                title = m[1].strip() if len(m) > 1 and m[1] else None
            else:
                name = m.strip()
                title = None
            if name and len(name) > 2:
                signors.append({"name": name, "title": title})

    seen = set()
    unique_signors = []
    for s in signors:
        key = (s["name"].lower(), (s["title"] or "").lower())
        if key not in seen:
            seen.add(key)
            unique_signors.append(s)
    result["signors"] = unique_signors[:6]

    return result


def create_agreement_excel(info: dict) -> bytes:
    """Create an Excel file from extracted agreement info."""
    rows = []

    # Date
    rows.append({
        "Category": "Agreement Date",
        "Name / Value": info.get("agreement_date") or "Not found",
        "Address / Title": ""
    })

    # Parties
    if info.get("parties"):
        for i, p in enumerate(info["parties"], 1):
            rows.append({
                "Category": f"Party {i}",
                "Name / Value": p.get("name", ""),
                "Address / Title": p.get("address", "")
            })
    else:
        rows.append({
            "Category": "Party",
            "Name / Value": "Not found",
            "Address / Title": ""
        })

    # Signors
    if info.get("signors"):
        for i, s in enumerate(info["signors"], 1):
            rows.append({
                "Category": f"Signor {i}",
                "Name / Value": s.get("name", ""),
                "Address / Title": s.get("title") or ""
            })
    else:
        rows.append({
            "Category": "Signor",
            "Name / Value": "Not found",
            "Address / Title": ""
        })

    df = pd.DataFrame(rows)
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Agreement Info")
    buffer.seek(0)
    return buffer.getvalue()


# ======================
# Session State Init
# ======================
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "user_email" not in st.session_state:
    st.session_state.user_email = None
if "verification_code" not in st.session_state:
    st.session_state.verification_code = None
if "pending_email" not in st.session_state:
    st.session_state.pending_email = None
if "code_sent" not in st.session_state:
    st.session_state.code_sent = False
if "merged_pdf" not in st.session_state:
    st.session_state.merged_pdf = None
if "single_pdf_bytes" not in st.session_state:
    st.session_state.single_pdf_bytes = None
if "single_pdf_name" not in st.session_state:
    st.session_state.single_pdf_name = None
if "total_pages" not in st.session_state:
    st.session_state.total_pages = 0
if "converted_docx" not in st.session_state:
    st.session_state.converted_docx = None
if "extracted_text" not in st.session_state:
    st.session_state.extracted_text = None
if "agreement_info" not in st.session_state:
    st.session_state.agreement_info = None


# ======================
# Sidebar - Login + Support
# ======================
st.sidebar.header("🔐 Login with Email")

if not st.session_state.logged_in:
    if not st.session_state.code_sent:
        with st.sidebar.form("email_form"):
            email = st.text_input("Email address", placeholder="you@example.com")
            send_btn = st.form_submit_button(
                "Send Verification Code", use_container_width=True, type="primary"
            )
            if send_btn:
                email = email.strip().lower()
                if not email:
                    st.error("Please enter your email.")
                elif not is_valid_email(email):
                    st.error("Please enter a valid email address.")
                else:
                    code = generate_verification_code()
                    st.session_state.verification_code = code
                    st.session_state.pending_email = email
                    st.session_state.code_sent = True
                    st.rerun()
    else:
        st.sidebar.info(f"Code sent to:\n**{st.session_state.pending_email}**")
        st.sidebar.warning(f"🧪 Demo Code: **{st.session_state.verification_code}**")

        with st.sidebar.form("verify_form"):
            user_code = st.text_input("Enter 6-digit verification code", max_chars=6)
            col1, col2 = st.columns(2)
            with col1:
                verify_btn = st.form_submit_button(
                    "Verify & Login", use_container_width=True, type="primary"
                )
            with col2:
                back_btn = st.form_submit_button("← Back", use_container_width=True)

            if back_btn:
                st.session_state.code_sent = False
                st.session_state.verification_code = None
                st.session_state.pending_email = None
                st.rerun()

            if verify_btn:
                if user_code.strip() == st.session_state.verification_code:
                    st.session_state.logged_in = True
                    st.session_state.user_email = st.session_state.pending_email
                    st.session_state.code_sent = False
                    st.session_state.verification_code = None
                    st.session_state.pending_email = None
                    st.success("Login successful!")
                    st.rerun()
                else:
                    st.error("Incorrect verification code. Please try again.")
else:
    st.sidebar.success(f"Logged in as:\n**{st.session_state.user_email}**")
    if st.sidebar.button("Logout", use_container_width=True):
        st.session_state.logged_in = False
        st.session_state.user_email = None
        st.session_state.merged_pdf = None
        st.session_state.single_pdf_bytes = None
        st.session_state.single_pdf_name = None
        st.session_state.total_pages = 0
        st.session_state.converted_docx = None
        st.session_state.extracted_text = None
        st.session_state.agreement_info = None
        st.rerun()

# Support the Writer
st.sidebar.markdown("---")
st.sidebar.subheader("💖 Support the Writer")
st.sidebar.markdown("If this tool helped you, consider buying me a coffee!")

SUPPORT_URL = "https://buymeacoffee.com/itmasterngw"

st.sidebar.link_button(
    "☕ Buy Me a Coffee",
    SUPPORT_URL,
    use_container_width=True,
    type="primary"
)


# ======================
# Main App
# ======================
if not st.session_state.logged_in:
    st.title("📄 PDF Tools")
    st.info("👈 Please login with your email in the sidebar to continue.")
    st.stop()

st.title("📄 PDF Tools")
st.markdown(f"Welcome, **{st.session_state.user_email}**!")

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🔗 Combine Multiple PDFs",
    "📄 Select Pages & Download",
    "🔄 Rotate Pages",
    "📝 PDF to Word",
    "📋 Extract Agreement Info"
])


# =====================================================
# TAB 1: Combine Multiple PDFs
# =====================================================
with tab1:
    st.subheader("Combine Multiple PDFs into One")
    st.markdown("Upload several PDF files → Merge them → Download the combined file")

    uploaded_files = st.file_uploader(
        "Choose one or more PDF files",
        type=["pdf"],
        accept_multiple_files=True,
        key="multi_uploader"
    )

    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)} file(s) uploaded")
        st.markdown("**Files in merge order:**")
        for i, f in enumerate(uploaded_files, 1):
            st.write(f"{i}. `{f.name}` ({f.size / 1024:.1f} KB)")

        st.markdown("---")
        if st.button("🔗 Combine All PDFs", type="primary", use_container_width=True, key="merge_btn"):
            with st.spinner("Merging PDFs..."):
                writer = merge_pdfs(uploaded_files)
                if writer is not None:
                    buffer = io.BytesIO()
                    writer.write(buffer)
                    buffer.seek(0)
                    st.session_state.merged_pdf = buffer.getvalue()
                    st.success(f"✅ Successfully combined {len(uploaded_files)} PDFs!")
                else:
                    st.error("Failed to merge PDFs.")

    if st.session_state.merged_pdf is not None:
        st.markdown("---")
        st.subheader("Download Combined PDF")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_name = f"combined_{timestamp}.pdf"
        custom_name = st.text_input(
            "Custom file name (optional)",
            value=default_name,
            key="merge_custom_name"
        )
        if not custom_name.lower().endswith(".pdf"):
            custom_name += ".pdf"

        st.download_button(
            label="📥 Download Combined PDF",
            data=st.session_state.merged_pdf,
            file_name=custom_name,
            mime="application/pdf",
            use_container_width=True,
            type="primary",
            key="download_merged"
        )
        st.caption(f"File size: {len(st.session_state.merged_pdf) / 1024:.1f} KB")

        with st.expander("👁️ Preview & Edit Combined PDF", expanded=False):
            edited = preview_and_edit_pdf(st.session_state.merged_pdf, key="merged_edit")
            if edited is not None:
                st.session_state.merged_pdf = edited
                st.rerun()


# =====================================================
# TAB 2: Select Pages & Download
# =====================================================
with tab2:
    st.subheader("Select Pages from PDF & Download")
    st.markdown("Upload a PDF → Choose which pages you want → Download with your own file name")

    single_file = st.file_uploader(
        "Upload a PDF file",
        type=["pdf"],
        accept_multiple_files=False,
        key="single_uploader"
    )

    if single_file is not None:
        pdf_bytes = single_file.read()
        st.session_state.single_pdf_bytes = pdf_bytes
        st.session_state.single_pdf_name = single_file.name

        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            total_pages = len(reader.pages)
            st.session_state.total_pages = total_pages
        except Exception as e:
            st.error(f"Cannot read this PDF: {e}")
            st.stop()

        st.success(f"✅ Uploaded: **{single_file.name}**")
        st.info(f"This PDF has **{total_pages}** page(s)")

        with st.expander("👁️ Preview & Edit Uploaded PDF", expanded=False):
            edited = preview_and_edit_pdf(pdf_bytes, key="single_edit")
            if edited is not None:
                st.session_state.single_pdf_bytes = edited
                st.session_state.total_pages = len(PdfReader(io.BytesIO(edited)).pages)
                st.rerun()

        st.markdown("---")
        st.subheader("1️⃣ Select Pages")

        selection_mode = st.radio(
            "How do you want to select pages?",
            options=["Select specific pages", "Select page range", "Download all pages"],
            horizontal=True,
            key="select_mode"
        )

        selected_pages = []
        if selection_mode == "Select specific pages":
            page_options = list(range(1, total_pages + 1))
            selected_pages = st.multiselect(
                "Choose pages (you can select multiple)",
                options=page_options,
                default=[1] if total_pages >= 1 else [],
                key="select_multiselect"
            )
        elif selection_mode == "Select page range":
            col1, col2 = st.columns(2)
            with col1:
                start_page = st.number_input(
                    "From page", min_value=1, max_value=total_pages, value=1, key="select_start"
                )
            with col2:
                end_page = st.number_input(
                    "To page", min_value=1, max_value=total_pages, value=total_pages, key="select_end"
                )
            if start_page > end_page:
                st.warning("Start page cannot be greater than end page.")
            else:
                selected_pages = list(range(start_page, end_page + 1))
                st.write(f"Selected pages: **{start_page} to {end_page}** ({len(selected_pages)} pages)")
        else:
            selected_pages = list(range(1, total_pages + 1))
            st.write(f"All **{total_pages}** pages will be included.")

        st.markdown("---")
        st.subheader("2️⃣ Download Selected Pages")

        if selected_pages:
            original_name = single_file.name
            if original_name.lower().endswith(".pdf"):
                original_name = original_name[:-4]

            custom_filename = st.text_input(
                "Enter your preferred file name",
                value=f"{original_name}_selected",
                key="single_custom_name"
            )
            custom_filename = custom_filename.strip()
            if not custom_filename:
                custom_filename = "selected_pages"
            if not custom_filename.lower().endswith(".pdf"):
                custom_filename += ".pdf"

            try:
                with st.spinner("Preparing your PDF..."):
                    extracted_pdf = extract_pages(pdf_bytes, selected_pages)

                st.download_button(
                    label="📥 Download Selected Pages",
                    data=extracted_pdf,
                    file_name=custom_filename,
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary",
                    key="download_selected"
                )
                st.success(f"Ready to download **{len(selected_pages)}** page(s)")
                st.caption(f"File name: **{custom_filename}**")
                st.caption(f"File size: {len(extracted_pdf) / 1024:.1f} KB")

                with st.expander("👁️ Preview Selected Pages", expanded=False):
                    preview_pdf(extracted_pdf, key="selected_preview")
            except Exception as e:
                st.error(f"Error preparing PDF: {e}")
        else:
            st.warning("Please select at least one page.")


# =====================================================
# TAB 3: Rotate Pages
# =====================================================
with tab3:
    st.subheader("Rotate Pages in PDF")
    st.markdown("Upload a PDF → Select pages → Choose rotation angle → Download the rotated file")

    rotate_file = st.file_uploader(
        "Upload a PDF file to rotate",
        type=["pdf"],
        accept_multiple_files=False,
        key="rotate_uploader"
    )

    if rotate_file is not None:
        rotate_pdf_bytes = rotate_file.read()

        try:
            reader = PdfReader(io.BytesIO(rotate_pdf_bytes))
            total_pages = len(reader.pages)
        except Exception as e:
            st.error(f"Cannot read this PDF: {e}")
            st.stop()

        st.success(f"✅ Uploaded: **{rotate_file.name}**")
        st.info(f"This PDF has **{total_pages}** page(s)")

        with st.expander("👁️ Preview & Edit Uploaded PDF", expanded=False):
            edited = preview_and_edit_pdf(rotate_pdf_bytes, key="rotate_edit")
            if edited is not None:
                st.download_button(
                    "📥 Download Edited PDF",
                    data=edited,
                    file_name="edited.pdf",
                    mime="application/pdf",
                    key="download_edited_rotate"
                )

        st.markdown("---")
        st.subheader("1️⃣ Select Pages to Rotate")

        rotate_mode = st.radio(
            "How do you want to select pages?",
            options=["Select specific pages", "Select page range", "Rotate all pages"],
            horizontal=True,
            key="rotate_mode"
        )

        pages_to_rotate = []

        if rotate_mode == "Select specific pages":
            page_options = list(range(1, total_pages + 1))
            pages_to_rotate = st.multiselect(
                "Choose pages to rotate",
                options=page_options,
                default=[1] if total_pages >= 1 else [],
                key="rotate_multiselect"
            )
        elif rotate_mode == "Select page range":
            col1, col2 = st.columns(2)
            with col1:
                start_page = st.number_input(
                    "From page", min_value=1, max_value=total_pages, value=1, key="rotate_start"
                )
            with col2:
                end_page = st.number_input(
                    "To page", min_value=1, max_value=total_pages, value=total_pages, key="rotate_end"
                )
            if start_page > end_page:
                st.warning("Start page cannot be greater than end page.")
            else:
                pages_to_rotate = list(range(start_page, end_page + 1))
                st.write(f"Pages to rotate: **{start_page} to {end_page}** ({len(pages_to_rotate)} pages)")
        else:
            pages_to_rotate = list(range(1, total_pages + 1))
            st.write(f"All **{total_pages}** pages will be rotated.")

        st.markdown("---")
        st.subheader("2️⃣ Choose Rotation Angle")

        angle = st.selectbox(
            "Rotate clockwise by",
            options=[90, 180, 270],
            format_func=lambda x: f"{x}° clockwise",
            key="rotate_angle"
        )

        st.markdown("---")
        st.subheader("3️⃣ Download Rotated PDF")

        if pages_to_rotate:
            original_name = rotate_file.name
            if original_name.lower().endswith(".pdf"):
                original_name = original_name[:-4]

            custom_filename = st.text_input(
                "Enter your preferred file name",
                value=f"{original_name}_rotated_{angle}",
                key="rotate_custom_name"
            )
            custom_filename = custom_filename.strip()
            if not custom_filename:
                custom_filename = f"rotated_{angle}"
            if not custom_filename.lower().endswith(".pdf"):
                custom_filename += ".pdf"

            try:
                with st.spinner("Rotating pages..."):
                    rotated_pdf = rotate_pages(rotate_pdf_bytes, pages_to_rotate, angle)

                st.download_button(
                    label="📥 Download Rotated PDF",
                    data=rotated_pdf,
                    file_name=custom_filename,
                    mime="application/pdf",
                    use_container_width=True,
                    type="primary",
                    key="download_rotated"
                )
                st.success(
                    f"Ready! **{len(pages_to_rotate)}** page(s) rotated by **{angle}°** clockwise"
                )
                st.caption(f"File name: **{custom_filename}**")
                st.caption(f"File size: {len(rotated_pdf) / 1024:.1f} KB")

                with st.expander("👁️ Preview Rotated PDF", expanded=False):
                    preview_pdf(rotated_pdf, key="rotated_preview")
            except Exception as e:
                st.error(f"Error rotating PDF: {e}")
        else:
            st.warning("Please select at least one page to rotate.")


# =====================================================
# TAB 4: PDF to Word
# =====================================================
with tab4:
    st.subheader("Convert PDF to Word (.docx)")
    st.markdown(
        "Upload a PDF → Convert it → **Edit the text** → Download editable Word file"
    )

    pdf_to_word_file = st.file_uploader(
        "Upload a PDF file to convert",
        type=["pdf"],
        accept_multiple_files=False,
        key="pdf2word_uploader"
    )

    if pdf_to_word_file is not None:
        pdf_bytes = pdf_to_word_file.read()
        st.success(f"✅ Uploaded: **{pdf_to_word_file.name}**")
        st.caption(f"File size: {len(pdf_bytes) / 1024:.1f} KB")

        is_scanned = not has_extractable_text(pdf_bytes)
        if is_scanned:
            st.warning(
                "⚠️ This looks like a **scanned / image-only PDF**. "
                "OCR will be used to extract real editable text."
            )
        else:
            st.info("✅ This PDF already contains selectable text.")

        with st.expander("👁️ Preview PDF", expanded=False):
            preview_pdf(pdf_bytes, key="pdf2word_preview")

        st.markdown("---")

        col1, col2 = st.columns(2)
        with col1:
            force_ocr = st.checkbox(
                "Force OCR (even if text is present)",
                value=False
            )
        with col2:
            ocr_lang = st.selectbox(
                "OCR Language",
                options=["eng", "chi_sim", "chi_tra", "jpn", "kor", "fra", "deu", "spa"],
                index=0
            )

        if st.button("📝 Convert / Extract Text", type="primary", use_container_width=True, key="convert_btn"):
            with st.spinner(
                "Extracting text... "
                + ("(using OCR – this may take longer)" if (force_ocr or is_scanned) else "")
            ):
                try:
                    if force_ocr or is_scanned:
                        import pypdfium2 as pdfium
                        import pytesseract

                        pdf = pdfium.PdfDocument(pdf_bytes)
                        all_text = []
                        for i in range(len(pdf)):
                            page = pdf[i]
                            bitmap = page.render(scale=2.0)
                            pil_image = bitmap.to_pil()
                            text = pytesseract.image_to_string(pil_image, lang=ocr_lang)
                            all_text.append(f"--- Page {i + 1} ---\n{text.strip()}")
                        extracted_text = "\n\n".join(all_text)
                    else:
                        reader = PdfReader(io.BytesIO(pdf_bytes))
                        pages_text = []
                        for i, page in enumerate(reader.pages):
                            text = page.extract_text() or ""
                            pages_text.append(f"--- Page {i + 1} ---\n{text.strip()}")
                        extracted_text = "\n\n".join(pages_text)

                    st.session_state.extracted_text = extracted_text
                    st.session_state.converted_docx = None
                    st.success("✅ Text extracted! You can now edit it below.")
                except Exception as e:
                    st.error(f"Extraction failed: {e}")

        if st.session_state.extracted_text:
            st.markdown("---")
            st.subheader("✏️ Edit the extracted text")

            edited_text = st.text_area(
                "You can freely edit the text below before generating the Word file:",
                value=st.session_state.extracted_text,
                height=400,
                key="edit_text_area"
            )
            st.session_state.extracted_text = edited_text

            if st.button("📄 Generate Word from edited text", type="primary", use_container_width=True):
                from docx import Document
                from docx.shared import Pt

                doc = Document()
                style = doc.styles["Normal"]
                font = style.font
                font.name = "Arial"
                font.size = Pt(11)

                pages = edited_text.split("--- Page ")
                for i, page_content in enumerate(pages):
                    if not page_content.strip():
                        continue
                    if i > 0:
                        doc.add_page_break()
                    content = page_content
                    if "---" in content[:30]:
                        content = content.split("---", 1)[-1].lstrip()
                    doc.add_paragraph(content.strip())

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.session_state.converted_docx = buffer.getvalue()
                st.success("✅ Word document generated from your edited text!")

        if st.session_state.converted_docx is not None:
            st.markdown("---")
            st.subheader("Download Word File")

            original_name = pdf_to_word_file.name
            if original_name.lower().endswith(".pdf"):
                original_name = original_name[:-4]

            custom_name = st.text_input(
                "Custom file name (optional)",
                value=f"{original_name}.docx",
                key="docx_custom_name"
            )
            if not custom_name.lower().endswith(".docx"):
                custom_name += ".docx"

            st.download_button(
                label="📥 Download Word Document",
                data=st.session_state.converted_docx,
                file_name=custom_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary",
                key="download_docx"
            )
            st.caption(f"File size: {len(st.session_state.converted_docx) / 1024:.1f} KB")


# =====================================================
# TAB 5: Extract Agreement Info
# =====================================================
with tab5:
    st.subheader("Extract Agreement Information")
    st.markdown(
        "Upload an agreement / contract PDF → Extract **Date**, **Party Name & Address**, and **Signor Name & Title**"
    )

    agreement_file = st.file_uploader(
        "Upload Agreement PDF",
        type=["pdf"],
        accept_multiple_files=False,
        key="agreement_uploader"
    )

    if agreement_file is not None:
        pdf_bytes = agreement_file.read()
        st.success(f"✅ Uploaded: **{agreement_file.name}**")

        use_ocr = st.checkbox("Force OCR (for scanned PDFs)", value=False, key="agreement_ocr")
        ocr_lang = st.selectbox(
            "OCR Language",
            options=["eng", "chi_sim", "chi_tra", "jpn", "kor", "fra", "deu", "spa"],
            index=0,
            key="agreement_lang"
        )

        if st.button("🔍 Extract Information", type="primary", use_container_width=True, key="extract_btn"):
            with st.spinner("Extracting information..."):
                try:
                    info = extract_agreement_info(pdf_bytes, use_ocr=use_ocr, ocr_lang=ocr_lang)
                    st.session_state.agreement_info = info
                    st.success("✅ Extraction completed!")
                except Exception as e:
                    st.error(f"Extraction failed: {e}")
                    st.session_state.agreement_info = None

        if st.session_state.agreement_info:
            info = st.session_state.agreement_info

            st.markdown("---")
            st.subheader("Extracted Results")

            # Agreement Date
            st.markdown("#### 📅 Agreement / Effective Date")
            if info["agreement_date"]:
                st.success(info["agreement_date"])
            else:
                st.warning("Not found automatically")

            # Parties
            st.markdown("#### 👥 Parties (Name & Address)")
            if info["parties"]:
                for i, p in enumerate(info["parties"], 1):
                    st.write(f"**Party {i}**")
                    st.write(f"- **Name:** {p['name']}")
                    st.write(f"- **Address:** {p['address'] or 'Not found'}")
            else:
                st.warning("No clear party information found")

            # Signors
            st.markdown("#### ✍️ Signor(s) Name & Title")
            if info["signors"]:
                for i, s in enumerate(info["signors"], 1):
                    st.write(f"**Signor {i}**")
                    st.write(f"- **Name:** {s['name']}")
                    st.write(f"- **Title:** {s['title'] or 'Not found'}")
            else:
                st.warning("No clear signor information found")

            # Download as Excel
            st.markdown("---")
            excel_bytes = create_agreement_excel(info)
            st.download_button(
                label="📥 Download as Excel",
                data=excel_bytes,
                file_name=f"agreement_info_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                type="primary",
                key="download_agreement_excel"
            )

            with st.expander("📄 View extracted raw text (for verification)"):
                st.text(info["raw_text_preview"])
